"""
Document Ingestion Pipeline for Telecom GenAI Platform.

Handles ingestion of structured and unstructured documents into ChromaDB
with telecom-optimized chunking strategies.
"""

import argparse
import asyncio
import hashlib
import json
import logging
import os
import re
import xml.etree.ElementTree as ET
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any

import chromadb
from chromadb.config import Settings
from langchain.text_splitter import (
    MarkdownTextSplitter,
    RecursiveCharacterTextSplitter,
)
from tqdm import tqdm

from src.rag_pipeline.embeddings import EmbeddingService, EmbeddingSettings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""

    CSV = "csv"
    JSON = "json"
    XML = "xml"
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    CONFIG = "config"  # Network configuration files
    SLA = "sla"  # Service Level Agreement documents


@dataclass
class Document:
    """Represents a loaded document with metadata."""

    content: str
    source_path: str
    document_type: DocumentType
    ingestion_timestamp: datetime = field(default_factory=datetime.utcnow)
    version: str = "1.0"
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def document_id(self) -> str:
        """Generate unique document ID based on content hash."""
        return hashlib.sha256(
            f"{self.source_path}:{self.content[:1000]}".encode()
        ).hexdigest()[:16]


@dataclass
class Chunk:
    """Represents a document chunk ready for embedding."""

    content: str
    document_id: str
    chunk_index: int
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def chunk_id(self) -> str:
        """Generate unique chunk ID."""
        return f"{self.document_id}_{self.chunk_index}"


class ChunkingStrategy(ABC):
    """Base class for chunking strategies."""

    @abstractmethod
    def chunk(self, document: Document) -> list[Chunk]:
        """Split document into chunks."""
        pass


class ProseChunker(ChunkingStrategy):
    """Chunking strategy for prose documents like runbooks and manuals."""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: list[str] | None = None,
    ):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators or ["\n\n", "\n", ". ", " ", ""],
            length_function=len,
        )

    def chunk(self, document: Document) -> list[Chunk]:
        texts = self.splitter.split_text(document.content)
        return [
            Chunk(
                content=text,
                document_id=document.document_id,
                chunk_index=i,
                metadata={
                    **document.metadata,
                    "source_path": document.source_path,
                    "document_type": document.document_type.value,
                    "chunk_strategy": "prose",
                },
            )
            for i, text in enumerate(texts)
        ]


class MarkdownChunker(ChunkingStrategy):
    """Chunking strategy for markdown documents preserving structure."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        self.splitter = MarkdownTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

    def chunk(self, document: Document) -> list[Chunk]:
        texts = self.splitter.split_text(document.content)
        chunks = []

        for i, text in enumerate(texts):
            # Extract section header if present
            section_header = self._extract_section_header(text)
            metadata = {
                **document.metadata,
                "source_path": document.source_path,
                "document_type": document.document_type.value,
                "chunk_strategy": "markdown",
            }
            if section_header:
                metadata["section_header"] = section_header

            chunks.append(
                Chunk(
                    content=text,
                    document_id=document.document_id,
                    chunk_index=i,
                    metadata=metadata,
                )
            )

        return chunks

    def _extract_section_header(self, text: str) -> str | None:
        """Extract the first markdown header from text."""
        match = re.search(r"^#{1,6}\s+(.+)$", text, re.MULTILINE)
        return match.group(1).strip() if match else None


class SemanticChunker(ChunkingStrategy):
    """
    Semantic chunking strategy for SLA and structured documents.

    Groups content by semantic boundaries (sections, clauses, definitions)
    rather than fixed character counts.
    """

    # Patterns for SLA document sections
    SECTION_PATTERNS = [
        r"(?:^|\n)((?:\d+\.)+\s+[A-Z][^\n]+)",  # Numbered sections (1.1 Title)
        r"(?:^|\n)((?:Article|Section|Clause)\s+\d+[:\.\s][^\n]+)",  # Article/Section headers
        r"(?:^|\n)((?:APPENDIX|SCHEDULE|EXHIBIT)\s+[A-Z0-9]+[:\.\s][^\n]*)",  # Appendices
        r"(?:^|\n)((?:Definition|Term)[s]?:?\s*\n)",  # Definition sections
    ]

    def __init__(
        self,
        min_chunk_size: int = 200,
        max_chunk_size: int = 2000,
        overlap_sentences: int = 1,
    ):
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap_sentences = overlap_sentences

    def chunk(self, document: Document) -> list[Chunk]:
        content = document.content
        sections = self._split_by_sections(content)

        chunks = []
        for section_title, section_content in sections:
            # Further split large sections
            if len(section_content) > self.max_chunk_size:
                sub_chunks = self._split_large_section(section_content)
            else:
                sub_chunks = [section_content] if len(section_content) >= self.min_chunk_size else []

            for sub_content in sub_chunks:
                metadata = {
                    **document.metadata,
                    "source_path": document.source_path,
                    "document_type": document.document_type.value,
                    "chunk_strategy": "semantic",
                }
                if section_title:
                    metadata["section_header"] = section_title.strip()

                chunks.append(
                    Chunk(
                        content=sub_content,
                        document_id=document.document_id,
                        chunk_index=len(chunks),
                        metadata=metadata,
                    )
                )

        # Fallback if no sections found
        if not chunks:
            fallback = ProseChunker(chunk_size=1000, chunk_overlap=200)
            return fallback.chunk(document)

        return chunks

    def _split_by_sections(self, content: str) -> list[tuple[str | None, str]]:
        """Split content by semantic section boundaries."""
        sections: list[tuple[str | None, str]] = []
        combined_pattern = "|".join(f"({p})" for p in self.SECTION_PATTERNS)

        # Find all section headers
        matches = list(re.finditer(combined_pattern, content, re.MULTILINE | re.IGNORECASE))

        if not matches:
            return [(None, content)]

        # Extract sections between headers
        for i, match in enumerate(matches):
            header = match.group(0).strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            section_content = content[start:end].strip()

            if section_content:
                sections.append((header, section_content))

        # Include content before first header if substantial
        first_match_start = matches[0].start()
        if first_match_start > self.min_chunk_size:
            preamble = content[:first_match_start].strip()
            if preamble:
                sections.insert(0, (None, preamble))

        return sections

    def _split_large_section(self, content: str) -> list[str]:
        """Split large sections by sentences with overlap."""
        # Split by sentence boundaries
        sentences = re.split(r"(?<=[.!?])\s+", content)
        chunks = []
        current_chunk: list[str] = []
        current_size = 0

        for sentence in sentences:
            sentence_size = len(sentence)

            if current_size + sentence_size > self.max_chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                # Keep overlap sentences
                current_chunk = current_chunk[-self.overlap_sentences:] if self.overlap_sentences else []
                current_size = sum(len(s) for s in current_chunk)

            current_chunk.append(sentence)
            current_size += sentence_size

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks


class NetworkConfigChunker(ChunkingStrategy):
    """
    Telecom-optimized chunking for network configuration files.

    Preserves configuration blocks (interfaces, routing, policies) as atomic units.
    """

    # Common telecom config block patterns
    BLOCK_PATTERNS = [
        r"(interface\s+[\w/.-]+[\s\S]*?)(?=\ninterface\s|\n!|\Z)",  # Interface configs
        r"(router\s+\w+[\s\S]*?)(?=\nrouter\s|\n!|\Z)",  # Routing protocols
        r"(policy-map\s+[\w-]+[\s\S]*?)(?=\npolicy-map\s|\n!|\Z)",  # QoS policies
        r"(class-map\s+[\w-]+[\s\S]*?)(?=\nclass-map\s|\n!|\Z)",  # Traffic classes
        r"(vlan\s+\d+[\s\S]*?)(?=\nvlan\s|\n!|\Z)",  # VLAN configs
        r"(access-list\s+[\w-]+[\s\S]*?)(?=\naccess-list\s|\n!|\Z)",  # ACLs
        r"(\[[\w:.-]+\][\s\S]*?)(?=\n\[|\Z)",  # INI-style sections
    ]

    def __init__(self, max_chunk_size: int = 2000):
        self.max_chunk_size = max_chunk_size

    def chunk(self, document: Document) -> list[Chunk]:
        content = document.content
        chunks = []
        seen_content = set()

        # Try to extract config blocks
        for pattern in self.BLOCK_PATTERNS:
            matches = re.findall(pattern, content, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                block = match.strip()
                if block and block not in seen_content:
                    seen_content.add(block)
                    # Split large blocks if needed
                    if len(block) > self.max_chunk_size:
                        sub_chunks = self._split_large_block(block)
                        for sub_chunk in sub_chunks:
                            chunks.append(self._create_chunk(document, sub_chunk, len(chunks)))
                    else:
                        chunks.append(self._create_chunk(document, block, len(chunks)))

        # If no blocks found, fall back to line-based chunking
        if not chunks:
            chunks = self._fallback_chunk(document)

        return chunks

    def _split_large_block(self, block: str) -> list[str]:
        """Split large config blocks while preserving line integrity."""
        lines = block.split("\n")
        sub_chunks = []
        current_chunk: list[str] = []
        current_size = 0

        for line in lines:
            line_size = len(line) + 1
            if current_size + line_size > self.max_chunk_size and current_chunk:
                sub_chunks.append("\n".join(current_chunk))
                current_chunk = [line]
                current_size = line_size
            else:
                current_chunk.append(line)
                current_size += line_size

        if current_chunk:
            sub_chunks.append("\n".join(current_chunk))

        return sub_chunks

    def _fallback_chunk(self, document: Document) -> list[Chunk]:
        """Fallback chunking for unrecognized config formats."""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.max_chunk_size,
            chunk_overlap=100,
            separators=["\n\n", "\n!", "\n", " "],
        )
        texts = splitter.split_text(document.content)
        return [
            self._create_chunk(document, text, i) for i, text in enumerate(texts)
        ]

    def _create_chunk(self, document: Document, content: str, index: int) -> Chunk:
        """Create a chunk with config-specific metadata."""
        # Try to identify config block type
        block_type = self._identify_block_type(content)

        return Chunk(
            content=content,
            document_id=document.document_id,
            chunk_index=index,
            metadata={
                **document.metadata,
                "source_path": document.source_path,
                "document_type": document.document_type.value,
                "chunk_strategy": "network_config",
                "config_block_type": block_type,
            },
        )

    def _identify_block_type(self, content: str) -> str:
        """Identify the type of configuration block."""
        first_line = content.split("\n")[0].lower().strip()

        if first_line.startswith("interface"):
            return "interface"
        elif first_line.startswith("router"):
            return "routing"
        elif first_line.startswith("policy-map"):
            return "qos_policy"
        elif first_line.startswith("class-map"):
            return "traffic_class"
        elif first_line.startswith("vlan"):
            return "vlan"
        elif first_line.startswith("access-list"):
            return "acl"
        elif first_line.startswith("["):
            return "ini_section"
        else:
            return "general"


class DocumentIngester:
    """
    Main document ingestion class.

    Loads documents from multiple sources (structured and unstructured),
    tracks source metadata, and prepares documents for chunking.
    """

    EXTENSION_MAP = {
        ".csv": DocumentType.CSV,
        ".json": DocumentType.JSON,
        ".xml": DocumentType.XML,
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".txt": DocumentType.TXT,
        ".md": DocumentType.MD,
        ".conf": DocumentType.CONFIG,
        ".cfg": DocumentType.CONFIG,
        ".config": DocumentType.CONFIG,
        ".sla": DocumentType.SLA,
    }

    # Filename patterns that indicate SLA documents (case-insensitive)
    SLA_FILENAME_PATTERNS = [
        r"sla[_\-\s]",
        r"service[_\-\s]level",
        r"agreement",
        r"contract",
    ]

    def __init__(self):
        self._loaders: dict[DocumentType, callable] = {
            DocumentType.CSV: self._load_csv,
            DocumentType.JSON: self._load_json,
            DocumentType.XML: self._load_xml,
            DocumentType.PDF: self._load_pdf,
            DocumentType.DOCX: self._load_docx,
            DocumentType.TXT: self._load_text,
            DocumentType.MD: self._load_text,
            DocumentType.CONFIG: self._load_text,
            DocumentType.SLA: self._load_text,  # SLA docs use text loader
        }

    def load_documents(self, source_path: str) -> list[Document]:
        """
        Load documents from a file or directory.

        Args:
            source_path: Path to file or directory containing documents.

        Returns:
            List of loaded Document objects.
        """
        path = Path(source_path)
        documents: list[Document] = []

        if path.is_file():
            doc = self._load_single_file(path)
            if doc:
                documents.append(doc)
        elif path.is_dir():
            for file_path in path.rglob("*"):
                if file_path.is_file():
                    doc = self._load_single_file(file_path)
                    if doc:
                        documents.append(doc)
        else:
            raise FileNotFoundError(f"Source path not found: {source_path}")

        logger.info(f"Loaded {len(documents)} documents from {source_path}")
        return documents

    def _load_single_file(self, file_path: Path) -> Document | None:
        """Load a single file and return a Document object."""
        suffix = file_path.suffix.lower()
        doc_type = self.EXTENSION_MAP.get(suffix)

        if not doc_type:
            logger.warning(f"Unsupported file type: {file_path}")
            return None

        # Check if filename indicates SLA document (override type for chunking)
        is_sla = self._is_sla_document(file_path)

        loader = self._loaders.get(doc_type)
        if not loader:
            logger.warning(f"No loader for document type: {doc_type}")
            return None

        try:
            content, metadata = loader(file_path)
            # Mark as SLA for semantic chunking if filename matches
            final_doc_type = DocumentType.SLA if is_sla else doc_type
            metadata["detected_as_sla"] = is_sla

            return Document(
                content=content,
                source_path=str(file_path),
                document_type=final_doc_type,
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to load {file_path}: {e}")
            return None

    def _is_sla_document(self, file_path: Path) -> bool:
        """Check if filename indicates an SLA document."""
        filename = file_path.stem.lower()
        return any(
            re.search(pattern, filename, re.IGNORECASE)
            for pattern in self.SLA_FILENAME_PATTERNS
        )

    def _load_csv(self, file_path: Path) -> tuple[str, dict]:
        """Load CSV file as text with metadata."""
        import csv

        content_lines = []
        metadata = {"row_count": 0, "columns": []}

        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            headers = next(reader, [])
            metadata["columns"] = headers

            content_lines.append(" | ".join(headers))
            content_lines.append("-" * 40)

            for row in reader:
                content_lines.append(" | ".join(row))
                metadata["row_count"] += 1

        return "\n".join(content_lines), metadata

    def _load_json(self, file_path: Path) -> tuple[str, dict]:
        """Load JSON file with structure preservation."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Pretty print for readability while maintaining structure
        content = json.dumps(data, indent=2)
        metadata = {
            "json_type": type(data).__name__,
            "top_level_keys": list(data.keys()) if isinstance(data, dict) else None,
        }

        return content, metadata

    def _load_xml(self, file_path: Path) -> tuple[str, dict]:
        """Load XML file and convert to readable format."""
        tree = ET.parse(file_path)
        root = tree.getroot()

        # Convert to string representation
        content = ET.tostring(root, encoding="unicode", method="xml")

        metadata = {
            "root_tag": root.tag,
            "child_tags": list({child.tag for child in root}),
        }

        return content, metadata

    def _load_pdf(self, file_path: Path) -> tuple[str, dict]:
        """Load PDF file and extract text."""
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf is required for PDF support: pip install pypdf")

        text_content = []
        metadata = {"page_count": 0}

        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            metadata["page_count"] = len(reader.pages)

            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_content.append(f"--- Page {page_num} ---\n{text}")

        return "\n\n".join(text_content), metadata

    def _load_docx(self, file_path: Path) -> tuple[str, dict]:
        """Load DOCX file and extract text."""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX support: pip install python-docx")

        doc = docx.Document(file_path)
        paragraphs = []
        metadata = {"paragraph_count": 0, "has_tables": False}

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                metadata["paragraph_count"] += 1

        # Extract tables
        if doc.tables:
            metadata["has_tables"] = True
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_text.append(row_text)
                paragraphs.append("\n".join(table_text))

        return "\n\n".join(paragraphs), metadata

    def _load_text(self, file_path: Path) -> tuple[str, dict]:
        """Load plain text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        metadata = {
            "line_count": content.count("\n") + 1,
            "char_count": len(content),
        }

        # Try to detect region tags for telecom docs
        region_match = re.search(r"region[:\s]+(\w+)", content, re.IGNORECASE)
        if region_match:
            metadata["region"] = region_match.group(1)

        return content, metadata


class IngestionPipeline:
    """
    Complete ingestion pipeline from documents to ChromaDB.

    Orchestrates document loading, chunking, embedding, and storage.
    Uses async EmbeddingService for efficient embedding generation.
    """

    def __init__(
        self,
        chroma_host: str = "localhost",
        chroma_port: int = 8100,
        embedding_settings: EmbeddingSettings | None = None,
    ):
        self.ingester = DocumentIngester()
        self.embedding_settings = embedding_settings or EmbeddingSettings()
        self._embedding_service: EmbeddingService | None = None

        # Initialize ChromaDB client
        self.chroma_client = chromadb.HttpClient(
            host=chroma_host,
            port=chroma_port,
            settings=Settings(anonymized_telemetry=False),
        )

        # Chunking strategies by document type
        self.chunkers: dict[DocumentType, ChunkingStrategy] = {
            DocumentType.PDF: ProseChunker(chunk_size=1000, chunk_overlap=200),
            DocumentType.DOCX: ProseChunker(chunk_size=1000, chunk_overlap=200),
            DocumentType.TXT: ProseChunker(chunk_size=1000, chunk_overlap=200),
            DocumentType.MD: MarkdownChunker(chunk_size=1000, chunk_overlap=100),
            DocumentType.CONFIG: NetworkConfigChunker(max_chunk_size=2000),
            DocumentType.CSV: ProseChunker(chunk_size=500, chunk_overlap=50),
            DocumentType.JSON: ProseChunker(chunk_size=1500, chunk_overlap=100),
            DocumentType.XML: ProseChunker(chunk_size=1500, chunk_overlap=100),
            DocumentType.SLA: SemanticChunker(min_chunk_size=200, max_chunk_size=2000),
        }

    def load_documents(self, source_path: str) -> list[Document]:
        """Load documents from source path."""
        return self.ingester.load_documents(source_path)

    def chunk_documents(self, docs: list[Document]) -> list[Chunk]:
        """Chunk documents using appropriate strategies."""
        all_chunks: list[Chunk] = []

        for doc in docs:
            chunker = self.chunkers.get(doc.document_type, ProseChunker())
            chunks = chunker.chunk(doc)
            all_chunks.extend(chunks)
            logger.info(
                f"Chunked {doc.source_path} into {len(chunks)} chunks "
                f"using {chunker.__class__.__name__}"
            )

        logger.info(f"Total chunks created: {len(all_chunks)}")
        return all_chunks

    async def store_embeddings(
        self, chunks: list[Chunk], collection_name: str
    ) -> dict[str, int]:
        """
        Generate embeddings and store in ChromaDB with deduplication.

        Args:
            chunks: List of chunks to embed and store.
            collection_name: ChromaDB collection name.

        Returns:
            Dictionary with storage statistics (stored, skipped, updated).
        """
        stats = {"stored": 0, "skipped": 0, "updated": 0}

        if not chunks:
            logger.warning("No chunks to store")
            return stats

        # Get or create collection
        collection = self.chroma_client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )

        # Extract all IDs and check for existing entries (deduplication)
        all_ids = [chunk.chunk_id for chunk in chunks]
        existing_ids = set()

        try:
            # Check which IDs already exist in the collection
            existing = collection.get(ids=all_ids, include=[])
            existing_ids = set(existing["ids"]) if existing["ids"] else set()
        except Exception:
            # Collection might be empty or IDs don't exist
            pass

        # Filter chunks: new ones to add, existing ones to update
        new_chunks = []
        update_chunks = []

        for chunk in chunks:
            if chunk.chunk_id in existing_ids:
                update_chunks.append(chunk)
            else:
                new_chunks.append(chunk)

        logger.info(
            f"Deduplication: {len(new_chunks)} new, {len(update_chunks)} updates, "
            f"{len(chunks) - len(new_chunks) - len(update_chunks)} unchanged"
        )

        # Initialize embedding service
        async with EmbeddingService(self.embedding_settings) as embedding_service:
            # Process new chunks
            if new_chunks:
                texts = [chunk.content for chunk in new_chunks]
                ids = [chunk.chunk_id for chunk in new_chunks]
                metadatas = [chunk.metadata for chunk in new_chunks]

                logger.info(f"Generating embeddings for {len(texts)} new chunks...")
                embeddings = await embedding_service.embed_documents(texts, show_progress=True)

                # Store in ChromaDB with progress bar
                batch_size = 100
                with tqdm(total=len(new_chunks), desc="Storing new chunks", unit="chunk") as pbar:
                    for i in range(0, len(new_chunks), batch_size):
                        batch_end = min(i + batch_size, len(new_chunks))
                        collection.add(
                            ids=ids[i:batch_end],
                            embeddings=embeddings[i:batch_end],
                            documents=texts[i:batch_end],
                            metadatas=metadatas[i:batch_end],
                        )
                        pbar.update(batch_end - i)
                        stats["stored"] += batch_end - i

            # Process updates (upsert)
            if update_chunks:
                texts = [chunk.content for chunk in update_chunks]
                ids = [chunk.chunk_id for chunk in update_chunks]
                metadatas = [chunk.metadata for chunk in update_chunks]

                logger.info(f"Generating embeddings for {len(texts)} updated chunks...")
                embeddings = await embedding_service.embed_documents(texts, show_progress=True)

                # Upsert in ChromaDB with progress bar
                batch_size = 100
                with tqdm(total=len(update_chunks), desc="Updating chunks", unit="chunk") as pbar:
                    for i in range(0, len(update_chunks), batch_size):
                        batch_end = min(i + batch_size, len(update_chunks))
                        collection.upsert(
                            ids=ids[i:batch_end],
                            embeddings=embeddings[i:batch_end],
                            documents=texts[i:batch_end],
                            metadatas=metadatas[i:batch_end],
                        )
                        pbar.update(batch_end - i)
                        stats["updated"] += batch_end - i

        logger.info(
            f"Storage complete: {stats['stored']} stored, {stats['updated']} updated, "
            f"{stats['skipped']} skipped in collection '{collection_name}'"
        )
        return stats

    def clear_collection(self, collection_name: str) -> bool:
        """
        Delete a ChromaDB collection.

        Args:
            collection_name: Name of collection to delete.

        Returns:
            True if successful, False otherwise.
        """
        try:
            self.chroma_client.delete_collection(collection_name)
            logger.info(f"Deleted collection '{collection_name}'")
            return True
        except Exception as e:
            logger.error(f"Failed to delete collection '{collection_name}': {e}")
            return False

    async def run(self, source_path: str, collection_name: str) -> dict[str, Any]:
        """
        Run the complete ingestion pipeline.

        Args:
            source_path: Path to documents to ingest.
            collection_name: ChromaDB collection name.

        Returns:
            Dictionary with ingestion statistics.
        """
        start_time = datetime.utcnow()

        # Load documents with progress bar
        logger.info(f"Loading documents from {source_path}...")
        documents = self.load_documents(source_path)
        if not documents:
            return {"status": "error", "message": "No documents loaded"}

        # Chunk documents with progress bar
        logger.info("Chunking documents...")
        chunks = []
        with tqdm(documents, desc="Chunking documents", unit="doc") as pbar:
            for doc in pbar:
                chunker = self.chunkers.get(doc.document_type, ProseChunker())
                doc_chunks = chunker.chunk(doc)
                chunks.extend(doc_chunks)
                pbar.set_postfix(chunks=len(chunks))

        if not chunks:
            return {"status": "error", "message": "No chunks created"}

        logger.info(f"Total chunks created: {len(chunks)}")

        # Store (async) - returns storage stats
        storage_stats = await self.store_embeddings(chunks, collection_name)

        end_time = datetime.utcnow()
        duration = (end_time - start_time).total_seconds()

        stats = {
            "status": "success",
            "documents_loaded": len(documents),
            "chunks_created": len(chunks),
            "chunks_stored": storage_stats["stored"],
            "chunks_updated": storage_stats["updated"],
            "chunks_skipped": storage_stats["skipped"],
            "collection_name": collection_name,
            "duration_seconds": duration,
            "document_types": {
                doc_type.value: sum(1 for d in documents if d.document_type == doc_type)
                for doc_type in DocumentType
                if any(d.document_type == doc_type for d in documents)
            },
        }

        logger.info(f"Ingestion complete: {stats}")
        return stats


async def async_main():
    """
    Async CLI entry point.

    Examples:
        # Ingest documents from a directory
        python -m src.rag_pipeline.ingest --source docs/sample_data/ --collection telecom_kb

        # Ingest a single file
        python -m src.rag_pipeline.ingest --source runbook.pdf --collection runbooks

        # Clear a collection before ingesting
        python -m src.rag_pipeline.ingest --clear-collection telecom_kb

        # Use local embeddings (no API key needed)
        python -m src.rag_pipeline.ingest --source docs/ --collection kb --embedding-provider sentence_transformers
    """
    parser = argparse.ArgumentParser(
        description="Ingest documents into Telecom GenAI Platform knowledge base",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python -m src.rag_pipeline.ingest --source docs/sample_data/ --collection telecom_kb
    python -m src.rag_pipeline.ingest --source runbook.pdf --collection runbooks --chunk-size 1500
    python -m src.rag_pipeline.ingest --source configs/ --collection network_configs --chroma-port 8100
    python -m src.rag_pipeline.ingest --clear-collection telecom_docs
        """,
    )

    parser.add_argument(
        "--source",
        help="Source file or directory containing documents to ingest",
    )
    parser.add_argument(
        "--collection",
        help="ChromaDB collection name to store embeddings",
    )
    parser.add_argument(
        "--clear-collection",
        metavar="NAME",
        help="Delete the specified ChromaDB collection and exit",
    )
    parser.add_argument(
        "--chroma-host",
        default=os.getenv("CHROMA_HOST", "localhost"),
        help="ChromaDB host (default: localhost)",
    )
    parser.add_argument(
        "--chroma-port",
        type=int,
        default=int(os.getenv("CHROMA_PORT", "8100")),
        help="ChromaDB port (default: 8100)",
    )
    parser.add_argument(
        "--embedding-provider",
        default=os.getenv("EMBEDDING_PROVIDER", "openai"),
        choices=["openai", "azure_openai", "sentence_transformers"],
        help="Embedding provider (default: openai)",
    )
    parser.add_argument(
        "--embedding-model",
        default=os.getenv("EMBEDDING_MODEL", "text-embedding-3-small"),
        help="Embedding model name (default: text-embedding-3-small)",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=1000,
        help="Default chunk size for prose documents (default: 1000)",
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=200,
        help="Chunk overlap for prose documents (default: 200)",
    )
    parser.add_argument(
        "--no-cache",
        action="store_true",
        help="Disable Redis embedding cache",
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose logging",
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Initialize pipeline (needed for both clear and ingest operations)
    pipeline = IngestionPipeline(
        chroma_host=args.chroma_host,
        chroma_port=args.chroma_port,
    )

    # Handle --clear-collection option
    if args.clear_collection:
        if pipeline.clear_collection(args.clear_collection):
            print(f"Successfully deleted collection '{args.clear_collection}'")
            exit(0)
        else:
            print(f"Failed to delete collection '{args.clear_collection}'")
            exit(1)

    # Validate required arguments for ingestion
    if not args.source:
        parser.error("--source is required for ingestion")
    if not args.collection:
        parser.error("--collection is required for ingestion")

    # Configure embedding settings
    from src.rag_pipeline.embeddings import EmbeddingProvider

    embedding_settings = EmbeddingSettings(
        provider=EmbeddingProvider(args.embedding_provider),
        model_name=args.embedding_model,
        cache_enabled=not args.no_cache,
    )
    pipeline.embedding_settings = embedding_settings

    # Update chunk sizes if specified
    if args.chunk_size != 1000 or args.chunk_overlap != 200:
        for doc_type in [DocumentType.PDF, DocumentType.DOCX, DocumentType.TXT]:
            pipeline.chunkers[doc_type] = ProseChunker(
                chunk_size=args.chunk_size,
                chunk_overlap=args.chunk_overlap,
            )

    # Run pipeline
    try:
        stats = await pipeline.run(args.source, args.collection)

        if stats["status"] == "success":
            print("\n" + "=" * 50)
            print("INGESTION COMPLETED SUCCESSFULLY")
            print("=" * 50)
            print(f"  Documents loaded:  {stats['documents_loaded']}")
            print(f"  Chunks created:    {stats['chunks_created']}")
            print(f"  Chunks stored:     {stats['chunks_stored']}")
            print(f"  Chunks updated:    {stats['chunks_updated']}")
            print(f"  Duplicates skipped:{stats['chunks_skipped']}")
            print(f"  Collection:        {stats['collection_name']}")
            print(f"  Duration:          {stats['duration_seconds']:.2f}s")
            print(f"  Document types:    {stats['document_types']}")
            print("=" * 50)
        else:
            print(f"\nIngestion failed: {stats.get('message', 'Unknown error')}")
            exit(1)

    except Exception as e:
        logger.exception("Ingestion failed")
        print(f"\nError: {e}")
        exit(1)


def main():
    """Synchronous CLI entry point."""
    asyncio.run(async_main())


if __name__ == "__main__":
    main()
